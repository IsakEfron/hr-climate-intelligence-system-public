import io
import base64
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT

def set_cell_background(cell, color_hex):
    """Pone color de fondo a una celda"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def configure_text(run, size_pt, bold=False, color_rgb=None, italic=False):
    """Ayuda rápida para configurar texto"""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

def generate_word_report(data):
    doc = Document()
    
    # 1. CONFIGURACIÓN DE PÁGINA (A4 Landscape con márgenes mínimos)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.7) 
    section.page_height = Inches(8.3)
    
    # Márgenes mínimos (0.3 pulgadas)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)

    # 2. TABLA MAESTRA (2 Columnas)
    main_table = doc.add_table(rows=1, cols=2)
    main_table.autofit = False
    main_table.allow_autofit = False
    
    # AJUSTE DE PROPORCIÓN: Más espacio a la izquierda para la imagen
    # Total disponible ~11.1 pulgadas
    main_table.columns[0].width = Inches(5)  # ~70% (Imagen grande)
    main_table.columns[1].width = Inches(2)  # ~30% (Datos compactos)
    
    left_cell = main_table.cell(0, 0)
    right_cell = main_table.cell(0, 1)
    
    # Quitar bordes tabla maestra
    for row in main_table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders {}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'.format(nsdecls('w')))
            tcPr.append(tcBorders)

    # --- COLUMNA IZQUIERDA (Gráfica) ---
    
    # Títulos
    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Centro de Trabajo {data['planta'].upper()}")
    configure_text(run, 22, True, RGBColor(30, 48, 110)) # Azul Canels

    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SATISFACCIÓN LABORAL GENERAL")
    configure_text(run, 14, True, RGBColor(0, 0, 0))

    p = left_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Periodo: {data['periodo']}")
    configure_text(run, 10, True, RGBColor(100, 100, 100))
    
    # Imagen
    if data['chart_image']:
        try:
            image_data = base64.b64decode(data['chart_image'].split(',')[1])
            image_stream = io.BytesIO(image_data)
            
            p = left_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            # IMAGEN GRANDE (7.5 pulgadas)
            run.add_picture(image_stream, width=Inches(5.2))
        except Exception as e:
            left_cell.add_paragraph(f"[Error imagen: {str(e)}]")

    # --- COLUMNA DERECHA (Datos) ---
    
    # 1. KPI Total Trabajadores (Cuadro Azul)
    t_kpi = right_cell.add_table(rows=1, cols=1)
    t_kpi.autofit = False
    t_kpi.columns[0].width = Inches(2.0) # Ancho compacto
    
    cell_kpi = t_kpi.cell(0, 0)
    set_cell_background(cell_kpi, "1E306E") # Azul Oscuro
    
    p = cell_kpi.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TOTAL DE TRABAJADORES")
    configure_text(run, 11, True, RGBColor(255, 255, 255))
    
    p = cell_kpi.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{data['total_trabajadores']}")
    configure_text(run, 28, True, RGBColor(255, 255, 255))
    
    # Filtros (Texto pequeño)
    p = right_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12) # Espacio antes de las tablas
    
    f = data.get('filtros', {})
    run = p.add_run("Filtros aplicados:\n")
    configure_text(run, 8, True, RGBColor(30, 48, 110))
    
    txt = f"Género: {f.get('genero')} | Antigüedad: {f.get('antiguedad')} | Área: {f.get('area')}"
    run = p.add_run(txt)
    configure_text(run, 8, False, RGBColor(80, 80, 80))

    # --- FUNCIÓN TABLAS COMPACTAS ---
    def create_compact_table(title, items):
        # Título (Sin salto de línea grande)
        p = right_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2) # Muy pegado a la tabla
        run = p.add_run(title)
        configure_text(run, 10, True, RGBColor(30, 48, 110))
        
        if not items:
            right_cell.add_paragraph("Sin datos").paragraph_format.space_after = Pt(8)
            return

        # Tabla (Ancho fijo 2.8 pulgadas)
        table = right_cell.add_table(rows=len(items) + 1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(2) # Categoría
        table.columns[1].width = Inches(0.6) # %
        
        # Bordes finos
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = parse_xml(r'<w:tcBorders {}><w:top w:val="single" w:sz="4" w:color="CCCCCC"/><w:left w:val="single" w:sz="4" w:color="CCCCCC"/><w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/><w:right w:val="single" w:sz="4" w:color="CCCCCC"/></w:tcBorders>'.format(nsdecls('w')))
                tcPr.append(tcBorders)

        # Encabezado
        h0 = table.cell(0, 0)
        h1 = table.cell(0, 1)
        set_cell_background(h0, "1E306E")
        set_cell_background(h1, "1E306E")
        
        # Texto Encabezado
        p = h0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Categoría")
        configure_text(r, 8, True, RGBColor(255, 255, 255))
        
        p = h1.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("%")
        configure_text(r, 8, True, RGBColor(255, 255, 255))

        # Filas
        for i, item in enumerate(items, start=1):
            c0 = table.cell(i, 0)
            c1 = table.cell(i, 1)
            set_cell_background(c0, "D9E1F2")
            set_cell_background(c1, "D9E1F2")
            
            p = c0.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(item[0]))
            configure_text(r, 8, False, RGBColor(0, 0, 0)) # Fuente 8 para ahorrar espacio
            
            p = c1.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(f"{item[1]}%")
            configure_text(r, 8, True, RGBColor(0, 0, 0))

        # Espacio reducido después de la tabla
        right_cell.add_paragraph().paragraph_format.space_after = Pt(10)

    # 3. Generar Tablas
    create_compact_table("ASPECTOS MEJOR EVALUADOS", data.get('top_3', []))
    create_compact_table("ASPECTOS CON MENOR PUNTUACIÓN", data.get('bottom_3', []))

    # Guardar
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream